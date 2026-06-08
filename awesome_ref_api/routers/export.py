import json
import os
import re
import uuid
import zipfile
from datetime import datetime, timezone
from io import BytesIO

from fastapi import APIRouter, Depends, Query, UploadFile, File
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy.orm import Session, joinedload

from database import get_db, utc_isoformat
from models import Reference, Note, User, Group, DailyPlan, DailyTask, StandaloneNote, NoteTag
from deps import get_current_user
from routers.references import _to_dict, _make_ref_key, _apply_fields

PDF_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "pdf_data")
NOTES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "note_mark_data")
IMAGES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "images_data")

MAX_IMPORT_SIZE = 100 * 1024 * 1024  # 100MB

router = APIRouter()

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def _group_refs(groups: list[dict], references: list[dict]) -> list[dict]:
    """Organise references under their groups. Returns ordered group dicts."""
    group_map = {g["id"]: {**g, "refs": []} for g in groups}
    ungrouped_key = None
    for g in groups:
        if g.get("name") == "未分组" or g.get("id") == "ungrouped":
            ungrouped_key = g["id"]
    for ref in references:
        placed = False
        for gid in ref.get("groupIds", []):
            if gid in group_map:
                group_map[gid]["refs"].append(ref)
                placed = True
        if not placed and ungrouped_key and ungrouped_key in group_map:
            group_map[ungrouped_key]["refs"].append(ref)
    return list(group_map.values())


def _find_cjk_font() -> str | None:
    """Return path to an available CJK-capable TrueType font or None."""
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyh.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    ]
    for p in candidates:
        if os.path.isfile(p):
            return p
    return None


# ---------------------------------------------------------------------------
# Markdown generator
# ---------------------------------------------------------------------------

def _build_markdown(groups: list[dict], references: list[dict],
                    notes: dict, exported_at: str) -> str:
    grouped = _group_refs(groups, references)
    lines: list[str] = []
    lines.append("# AwesomeRef 导出数据")
    lines.append("")
    lines.append(f"**导出时间**: {exported_at}")
    lines.append(f"**文献数**: {len(references)}　**分组数**: {len(groups)}")
    lines.append("")

    for g in grouped:
        lines.append(f"## {g['name']}")
        lines.append("")
        if not g["refs"]:
            lines.append("*(空分组)*")
            lines.append("")
            continue
        for i, ref in enumerate(g["refs"], 1):
            authors = "; ".join(ref.get("authors") or [])
            title = ref.get("title") or "无标题"
            lines.append(f"### {i}. {title}")
            lines.append("")
            if authors:
                lines.append(f"- **作者**: {authors}")
            if ref.get("year"):
                lines.append(f"- **年份**: {ref['year']}")
            if ref.get("type"):
                lines.append(f"- **类型**: {ref['type']}")
            if ref.get("journal"):
                lines.append(f"- **期刊**: {ref['journal']}")
            if ref.get("volume"):
                lines.append(f"- **卷**: {ref['volume']}")
            if ref.get("issue"):
                lines.append(f"- **期**: {ref['issue']}")
            if ref.get("pages"):
                lines.append(f"- **页码**: {ref['pages']}")
            if ref.get("doi"):
                lines.append(f"- **DOI**: [{ref['doi']}](https://doi.org/{ref['doi']})")
            if ref.get("keywords"):
                lines.append(f"- **关键词**: {', '.join(ref['keywords'])}")
            lines.append("")
            if ref.get("abstract"):
                lines.append(f"**摘要**: {ref['abstract']}")
                lines.append("")
            ref_key = ref.get("id", "")
            if ref_key in notes and notes[ref_key].get("content"):
                lines.append(f"**笔记**: {notes[ref_key]['content']}")
                lines.append("")
            lines.append("---")
            lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# PDF generator (fpdf2)
# ---------------------------------------------------------------------------

def _build_pdf(groups: list[dict], references: list[dict],
               notes: dict, exported_at: str) -> BytesIO:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_left_margin(12)
    pdf.set_right_margin(12)
    w = pdf.epw  # effective page width
    left_margin = pdf.l_margin

    cjk_path = _find_cjk_font()
    if cjk_path:
        pdf.add_font("cjk", "", cjk_path, uni=True)
        pdf.add_font("cjk", "B", cjk_path, uni=True)
        body_font = "cjk"
    else:
        body_font = "Helvetica"

    # ── header ──
    pdf.add_page()
    pdf.set_font(body_font, "B", 18)
    pdf.cell(w, 12, "AwesomeRef Export", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font(body_font, "", 9)
    pdf.cell(w, 6, f"Export time: {exported_at}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.cell(w, 6, f"References: {len(references)}    Groups: {len(groups)}",
             new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)

    grouped = _group_refs(groups, references)

    for g in grouped:
        # ── group heading ──
        pdf.set_font(body_font, "B", 13)
        pdf.cell(w, 8, g["name"], new_x="LMARGIN", new_y="NEXT")
        y = pdf.get_y()
        pdf.set_draw_color(180)
        pdf.line(left_margin, y + 1, left_margin + w, y + 1)
        pdf.set_draw_color(0)
        pdf.ln(5)

        if not g["refs"]:
            pdf.set_font(body_font, "", 9)
            pdf.cell(w, 6, "(empty)", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(4)
            continue

        for i, ref in enumerate(g["refs"], 1):
            # Check if we need a page break before this reference
            # (keep at least 25mm of space for the first few lines)
            if pdf.y > pdf.h - 30:
                pdf.add_page()

            # ── reference title ──
            title = ref.get("title") or "Untitled"
            pdf.set_font(body_font, "B", 10.5)
            pdf.multi_cell(w, 5.5, f"{i}. {title}", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font(body_font, "", 9)

            # ── metadata block (label: value on same line) ──
            lines = []
            authors = "; ".join(ref.get("authors") or [])
            if authors:
                lines.append(("Authors:", authors))

            meta = []
            if ref.get("year"):
                meta.append(ref["year"])
            if ref.get("journal"):
                meta.append(ref["journal"])
            if ref.get("volume"):
                meta.append(f"Vol.{ref['volume']}")
            if ref.get("issue"):
                meta.append(f"No.{ref['issue']}")
            if ref.get("pages"):
                meta.append(f"pp.{ref['pages']}")
            if meta:
                lines.append(("Source:", "  ".join(meta)))

            if ref.get("doi"):
                lines.append(("DOI:", ref["doi"]))
            if ref.get("keywords"):
                lines.append(("Keywords:", ", ".join(ref["keywords"])))

            for label, value in lines:
                pdf.set_font(body_font, "B", 9)
                label_w = pdf.get_string_width(label + " ")
                pdf.cell(label_w, 5, label + " ")
                pdf.set_font(body_font, "", 9)
                pdf.multi_cell(w - label_w, 5, value, new_x="LMARGIN", new_y="NEXT")

            # ── abstract ──
            if ref.get("abstract"):
                pdf.ln(1.5)
                pdf.set_font(body_font, "B", 8.5)
                pdf.cell(w, 4.5, "Abstract:", new_x="LMARGIN", new_y="NEXT")
                pdf.set_font(body_font, "", 8.5)
                pdf.multi_cell(w, 4.5, ref["abstract"], new_x="LMARGIN", new_y="NEXT")

            # ── note ──
            ref_key = ref.get("id", "")
            if ref_key in notes and notes[ref_key].get("content"):
                pdf.ln(1.5)
                pdf.set_font(body_font, "B", 8.5)
                pdf.cell(w, 4.5, "Notes:", new_x="LMARGIN", new_y="NEXT")
                pdf.set_font(body_font, "", 8.5)
                pdf.multi_cell(w, 4.5, notes[ref_key]["content"],
                               new_x="LMARGIN", new_y="NEXT")

            # ── separator between references ──
            pdf.ln(2.5)
            y = pdf.get_y()
            pdf.set_draw_color(210)
            pdf.line(left_margin, y, left_margin + w, y)
            pdf.set_draw_color(0)
            pdf.ln(3)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Word / docx generator (python-docx)
# ---------------------------------------------------------------------------

def _build_docx(groups: list[dict], references: list[dict],
                notes: dict, exported_at: str) -> BytesIO:
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)

    doc.add_heading("AwesomeRef Export", level=0)
    doc.add_paragraph(f"Export time: {exported_at}")
    doc.add_paragraph(f"References: {len(references)}    Groups: {len(groups)}")

    grouped = _group_refs(groups, references)
    for g in grouped:
        doc.add_heading(g["name"], level=1)
        if not g["refs"]:
            doc.add_paragraph("(empty)")
            continue
        for i, ref in enumerate(g["refs"], 1):
            doc.add_heading(f"{i}. {ref.get('title') or 'Untitled'}", level=2)
            authors = "; ".join(ref.get("authors") or [])
            if authors:
                doc.add_paragraph("Authors: ", style="List Bullet").add_run(authors)
            parts = []
            if ref.get("year"):
                parts.append(("Year", ref["year"]))
            if ref.get("type"):
                parts.append(("Type", ref["type"]))
            if ref.get("journal"):
                parts.append(("Journal", ref["journal"]))
            if ref.get("volume"):
                parts.append(("Volume", ref["volume"]))
            if ref.get("issue"):
                parts.append(("Issue", ref["issue"]))
            if ref.get("pages"):
                parts.append(("Pages", ref["pages"]))
            for label, val in parts:
                doc.add_paragraph(f"{label}: {val}", style="List Bullet")
            if ref.get("doi"):
                doc.add_paragraph(f"DOI: {ref['doi']}", style="List Bullet")
            if ref.get("keywords"):
                doc.add_paragraph("Keywords: ", style="List Bullet").add_run(
                    ", ".join(ref["keywords"]))
            if ref.get("abstract"):
                doc.add_paragraph("Abstract", style="List Bullet")
                doc.add_paragraph(ref["abstract"])
            ref_key = ref.get("id", "")
            if ref_key in notes and notes[ref_key].get("content"):
                p = doc.add_paragraph("Notes: ", style="List Bullet")
                p.add_run(notes[ref_key]["content"])
            doc.add_paragraph("─" * 40)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# ZIP backup generator
# ---------------------------------------------------------------------------

def _build_zip(user, db, groups, references, notes, daily_plans, exported_at) -> BytesIO:
    # Include trashed references
    trashed = [
        _to_dict(r)
        for r in db.query(Reference)
                    .options(joinedload(Reference.groups))
                    .filter(Reference.user_id == user.id,
                            Reference.deleted_at.isnot(None))
                    .all()
    ]
    all_refs = references + trashed

    # Standalone notes
    standalone_notes = []
    sn_rows = db.query(StandaloneNote).filter(StandaloneNote.user_id == user.id).all()
    for n in sn_rows:
        if not n.filename:
            continue
        filepath = os.path.join(NOTES_DIR, n.filename)
        content = ""
        if os.path.exists(filepath):
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
        tags = [{"name": t.name, "color": t.color} for t in n.tags]
        standalone_notes.append({
            "title": n.title,
            "content": content,
            "pinned": bool(n.pinned),
            "tags": tags,
            "createdAt": utc_isoformat(n.created_at),
            "updatedAt": utc_isoformat(n.updated_at),
        })

    # Note tags
    note_tags = [{"name": t.name, "color": t.color}
                 for t in db.query(NoteTag).filter(NoteTag.user_id == user.id).all()]

    data = {
        "export_version": "2.0",
        "exported_at": exported_at,
        "groups": groups,
        "references": all_refs,
        "notes": notes,
        "standalone_notes": standalone_notes,
        "note_tags": note_tags,
        "daily_plans": daily_plans,
    }

    buf = BytesIO()
    _note_count = 0
    _pdf_count = 0
    _pdf_missing = 0
    _img_count = 0
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("data.json", json.dumps(data, ensure_ascii=False, indent=2))

        # Collect standalone note .md files
        for n in sn_rows:
            if not n.filename:
                continue
            filepath = os.path.join(NOTES_DIR, n.filename)
            if os.path.exists(filepath):
                zf.write(filepath, f"notes/{n.filename}")
                _note_count += 1

        # Collect PDF files
        for ref in db.query(Reference).filter(Reference.user_id == user.id).all():
            if ref.pdf_filename:
                pdf_path = os.path.join(PDF_DIR, ref.pdf_filename)
                if os.path.exists(pdf_path):
                    zf.write(pdf_path, f"pdfs/{ref.pdf_filename}")
                    _pdf_count += 1
                else:
                    _pdf_missing += 1
                    print(f"[export_zip] PDF文件缺失: {ref.pdf_filename} (路径: {pdf_path})")

        # Collect images referenced in standalone notes
        image_pattern = r'/api/standalone-notes/images/([a-f0-9]+\.\w+)'
        seen_images = set()
        for n in sn_rows:
            if not n.filename:
                continue
            filepath = os.path.join(NOTES_DIR, n.filename)
            if not os.path.exists(filepath):
                continue
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
            for match in re.finditer(image_pattern, content):
                img_name = match.group(1)
                if img_name in seen_images:
                    continue
                seen_images.add(img_name)
                img_path = os.path.join(IMAGES_DIR, img_name)
                if os.path.exists(img_path):
                    zf.write(img_path, f"images/{img_name}")
                    _img_count += 1

    print(f"[export_zip] 完成: 笔记={_note_count}, PDF={_pdf_count}, PDF缺失={_pdf_missing}, 图片={_img_count}")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@router.get("/export")
def export_data(
    fmt: str = Query("json", alias="format"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    groups = [{"id": g.group_key, "name": g.name}
              for g in db.query(Group).filter(Group.user_id == user.id).all()]
    references = [
        _to_dict(r)
        for r in db.query(Reference)
                    .options(joinedload(Reference.groups))
                    .filter(Reference.user_id == user.id,
                            Reference.deleted_at.is_(None))
                    .all()
    ]
    notes = {n.ref_key: {"content": n.content}
             for n in db.query(Note).filter(Note.user_id == user.id).all()}

    daily_plans = []
    for plan in db.query(DailyPlan).filter(DailyPlan.user_id == user.id).all():
        tasks = sorted(plan.tasks, key=lambda t: (t.sort_order, t.id))
        daily_plans.append({
            "date": plan.date,
            "tasks": [{"title": t.title, "status": t.status, "note": t.note or "", "sortOrder": t.sort_order} for t in tasks],
        })

    exported_at = datetime.now(timezone.utc).isoformat()

    if fmt == "md":
        md_text = _build_markdown(groups, references, notes, exported_at)
        return Response(
            content=md_text.encode("utf-8"),
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition":
                     f"attachment; filename=awesomeref-export-{exported_at[:10]}.md"},
        )

    if fmt == "pdf":
        buf = _build_pdf(groups, references, notes, exported_at)
        return Response(
            content=buf.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition":
                     f"attachment; filename=awesomeref-export-{exported_at[:10]}.pdf"},
        )

    if fmt == "docx":
        buf = _build_docx(groups, references, notes, exported_at)
        return Response(
            content=buf.getvalue(),
            media_type="application/"
                        "vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition":
                     f"attachment; filename=awesomeref-export-{exported_at[:10]}.docx"},
        )

    if fmt == "zip":
        buf = _build_zip(user, db, groups, references, notes, daily_plans, exported_at)
        return Response(
            content=buf.getvalue(),
            media_type="application/zip",
            headers={"Content-Disposition":
                     f"attachment; filename=awesomeref-backup-{exported_at[:10]}.zip"},
        )

    # Default: json
    standalone_notes = []
    for n in db.query(StandaloneNote).filter(StandaloneNote.user_id == user.id).all():
        filepath = os.path.join(NOTES_DIR, n.filename) if n.filename else None
        content = ""
        if filepath and os.path.exists(filepath):
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
        tags = [{"name": t.name, "color": t.color} for t in n.tags]
        standalone_notes.append({
            "title": n.title,
            "content": content,
            "pinned": bool(n.pinned),
            "tags": tags,
            "createdAt": utc_isoformat(n.created_at),
            "updatedAt": utc_isoformat(n.updated_at),
        })

    note_tags = [{"name": t.name, "color": t.color}
                 for t in db.query(NoteTag).filter(NoteTag.user_id == user.id).all()]

    return {
        "export_version": "2.0",
        "exported_at": exported_at,
        "groups": groups,
        "references": references,
        "notes": notes,
        "standalone_notes": standalone_notes,
        "note_tags": note_tags,
        "daily_plans": daily_plans,
    }


class ImportData(BaseModel):
    export_version: str = ""
    exported_at: str = ""
    groups: list = []
    references: list = []
    notes: dict = {}
    standalone_notes: list = []
    note_tags: list = []
    daily_plans: list = []


# ---------------------------------------------------------------------------
# Shared import helpers
# ---------------------------------------------------------------------------

def _import_groups(user_id: int, groups_data: list, db: Session) -> dict:
    """Import groups, return old_key -> new_key mapping."""
    existing_groups = {g.name: g for g in db.query(Group).filter(Group.user_id == user_id).all()}
    group_key_map = {}
    for g in groups_data:
        name = (g.get("name") or "").strip()
        old_key = g.get("id", "")
        if not name:
            continue
        if name in existing_groups:
            group_key_map[old_key] = existing_groups[name].group_key
        else:
            new_key = f"grp-{uuid.uuid4().hex[:12]}"
            new_group = Group(user_id=user_id, group_key=new_key, name=name)
            db.add(new_group)
            db.flush()
            existing_groups[name] = new_group
            group_key_map[old_key] = new_key
    return group_key_map


def _ensure_ungrouped(user_id: int, db: Session) -> Group:
    ungrouped = db.query(Group).filter(Group.user_id == user_id, Group.group_key == "ungrouped").first()
    if not ungrouped:
        ungrouped = Group(user_id=user_id, group_key="ungrouped", name="未分组")
        db.add(ungrouped)
        db.flush()
    return ungrouped


def _import_references(user_id: int, refs_data: list, group_key_map: dict, db: Session) -> dict:
    """Import references, return old_ref_key -> new_ref_key mapping."""
    existing_refs = {r.ref_key: r for r in db.query(Reference).filter(Reference.user_id == user_id, Reference.deleted_at.is_(None)).all()}
    trashed_refs = {r.ref_key: r for r in db.query(Reference).filter(Reference.user_id == user_id, Reference.deleted_at.isnot(None)).all()}
    ungrouped = _ensure_ungrouped(user_id, db)
    ref_key_map = {}

    for item in refs_data:
        old_ref_key = item.get("id", "")
        title = (item.get("title") or "").strip()
        ref_key = _make_ref_key(title)
        if ref_key in existing_refs and existing_refs[ref_key].title.lower().strip() != title.lower().strip():
            ref_key = ref_key + uuid.uuid4().hex[:8]

        ref = existing_refs.get(ref_key)
        trashed_ref = trashed_refs.get(ref_key) if not ref else None

        if ref:
            _apply_fields(ref, item)
        elif trashed_ref:
            trashed_ref.deleted_at = None
            _apply_fields(trashed_ref, item)
            ref = trashed_ref
            existing_refs[ref_key] = ref
        else:
            ref = Reference(user_id=user_id, ref_key=ref_key)
            _apply_fields(ref, item)
            db.add(ref)
            db.flush()
            existing_refs[ref_key] = ref

        if old_ref_key:
            ref_key_map[old_ref_key] = ref_key

        raw_group_ids = item.get("groupIds", [])
        mapped_ids = [group_key_map.get(gid, gid) for gid in raw_group_ids]
        if raw_group_ids:
            ref.groups.clear()
        if mapped_ids:
            for gid in mapped_ids:
                g = db.query(Group).filter(Group.user_id == user_id, Group.group_key == gid).first()
                if g:
                    ref.groups.append(g)
        if not ref.groups:
            ref.groups.append(ungrouped)

    db.commit()
    return ref_key_map


def _import_notes(user_id: int, notes_data: dict, ref_key_map: dict, db: Session):
    """Import notes with merge logic (concatenation via separator)."""
    for old_ref_key, note_data in notes_data.items():
        content = note_data.get("content", "").strip()
        if not content:
            continue
        ref_key = ref_key_map.get(old_ref_key, old_ref_key)
        now = datetime.now(timezone.utc)
        note = db.query(Note).filter(Note.user_id == user_id, Note.ref_key == ref_key).first()
        if note:
            existing = (note.content or "").strip()
            if not existing:
                note.content = content
                note.updated_at = now
            elif existing == content:
                pass
            else:
                note.content = existing + "\n\n---\n\n" + content
                note.updated_at = now
        else:
            note = Note(user_id=user_id, ref_key=ref_key, content=content, updated_at=now)
            db.add(note)
    db.commit()


def _import_tags(user_id: int, tags_data: list, db: Session) -> dict:
    """Import note tags, return name -> NoteTag mapping."""
    existing_tags = {t.name: t for t in db.query(NoteTag).filter(NoteTag.user_id == user_id).all()}
    tag_map = dict(existing_tags)
    for tag_data in tags_data:
        name = (tag_data.get("name") or "").strip()
        if not name or name in tag_map:
            continue
        tag = NoteTag(user_id=user_id, name=name, color=tag_data.get("color", "#409eff"))
        db.add(tag)
        db.flush()
        tag_map[name] = tag
    db.commit()
    return tag_map


def _import_standalone_notes(user_id: int, sn_data_list: list, tag_map: dict, db: Session,
                             zip_file: zipfile.ZipFile | None = None):
    """Import standalone notes. If zip_file provided, tries to read content from it."""
    existing_sn = {n.title: n for n in db.query(StandaloneNote).filter(StandaloneNote.user_id == user_id).all()}
    os.makedirs(NOTES_DIR, exist_ok=True)

    for sn_data in sn_data_list:
        title = (sn_data.get("title") or "").strip()
        if not title or title in existing_sn:
            continue

        now = datetime.now(timezone.utc)
        sn = StandaloneNote(user_id=user_id, title=title, filename="", pinned=1 if sn_data.get("pinned") else 0,
                            created_at=now, updated_at=now)
        db.add(sn)
        db.flush()

        filename = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '', title).strip()
        filename = re.sub(r'\.{2,}', '', filename)
        if not filename:
            filename = "无标题笔记"
        filename = f"{filename}_{sn.id}.md"
        sn.filename = filename

        sn_tags = sn_data.get("tags", [])
        if sn_tags:
            tag_objs = [tag_map[t["name"]] for t in sn_tags if t.get("name") in tag_map]
            sn.tags = tag_objs

        note_content = sn_data.get("content", "")
        filepath = os.path.join(NOTES_DIR, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(note_content)

        existing_sn[title] = sn

    db.commit()


def _import_daily_plans(user_id: int, plans_data: list, db: Session):
    """Import daily plans with task deduplication."""
    for plan_data in plans_data:
        date = plan_data.get("date", "").strip()
        if not date:
            continue
        plan = db.query(DailyPlan).filter(DailyPlan.user_id == user_id, DailyPlan.date == date).first()
        if not plan:
            plan = DailyPlan(user_id=user_id, date=date)
            db.add(plan)
            db.flush()

        tasks_data = plan_data.get("tasks", [])
        existing_titles = {t.title for t in plan.tasks}
        max_order = max((t.sort_order for t in plan.tasks), default=-1)

        for td in tasks_data:
            td_title = td.get("title", "")
            if td_title in existing_titles:
                continue
            max_order += 1
            task = DailyTask(
                plan_id=plan.id,
                title=td_title,
                status=td.get("status", "pending"),
                note=td.get("note", ""),
                sort_order=td.get("sortOrder", max_order),
            )
            db.add(task)
    db.commit()


@router.post("/import")
def import_data(data: ImportData, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    group_key_map = _import_groups(user.id, data.groups, db)
    ref_key_map = _import_references(user.id, data.references, group_key_map, db)
    _import_notes(user.id, data.notes, ref_key_map, db)
    tag_map = _import_tags(user.id, data.note_tags, db)
    _import_standalone_notes(user.id, data.standalone_notes, tag_map, db)
    _import_daily_plans(user.id, data.daily_plans, db)
    return {"success": True}


@router.post("/import/zip")
async def import_zip(file: UploadFile = File(...), user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    content = await file.read()
    if len(content) > MAX_IMPORT_SIZE:
        return Response(status_code=413, content='{"detail":"文件大小超过 100MB 限制"}', media_type="application/json")
    buf = BytesIO(content)

    try:
        zf = zipfile.ZipFile(buf, "r")
    except zipfile.BadZipFile:
        return Response(status_code=400, content='{"detail":"无效的 ZIP 文件"}', media_type="application/json")

    if "data.json" not in zf.namelist():
        zf.close()
        return Response(status_code=400, content='{"detail":"ZIP 中缺少 data.json"}', media_type="application/json")

    data = json.loads(zf.read("data.json"))

    # --- Import groups ---
    group_key_map = _import_groups(user.id, data.get("groups", []), db)

    # --- Import references (with PDF restoration from ZIP) ---
    existing_refs = {r.ref_key: r for r in db.query(Reference).filter(Reference.user_id == user.id, Reference.deleted_at.is_(None)).all()}
    trashed_refs = {r.ref_key: r for r in db.query(Reference).filter(Reference.user_id == user.id, Reference.deleted_at.isnot(None)).all()}
    ungrouped = _ensure_ungrouped(user.id, db)
    ref_key_map = {}

    for item in data.get("references", []):
        old_ref_key = item.get("id", "")
        title = (item.get("title") or "").strip()
        ref_key = _make_ref_key(title)
        if ref_key in existing_refs and existing_refs[ref_key].title.lower().strip() != title.lower().strip():
            ref_key = ref_key + uuid.uuid4().hex[:8]

        ref = existing_refs.get(ref_key)
        trashed_ref = trashed_refs.get(ref_key) if not ref else None

        if ref:
            _apply_fields(ref, item)
        elif trashed_ref:
            trashed_ref.deleted_at = None
            _apply_fields(trashed_ref, item)
            ref = trashed_ref
            existing_refs[ref_key] = ref
        else:
            ref = Reference(user_id=user.id, ref_key=ref_key)
            _apply_fields(ref, item)
            db.add(ref)
            db.flush()
            existing_refs[ref_key] = ref

        if old_ref_key:
            ref_key_map[old_ref_key] = ref_key

        # Restore PDF file from ZIP (sanitized against path traversal)
        pdf_filename = item.get("pdfFilename")
        if pdf_filename:
            pdf_filename = os.path.basename(pdf_filename)
        if pdf_filename and f"pdfs/{pdf_filename}" in zf.namelist():
            os.makedirs(PDF_DIR, exist_ok=True)
            pdf_path = os.path.join(PDF_DIR, pdf_filename)
            if not os.path.exists(pdf_path):
                with open(pdf_path, "wb") as f:
                    f.write(zf.read(f"pdfs/{pdf_filename}"))
            ref.pdf_filename = pdf_filename

        raw_group_ids = item.get("groupIds", [])
        mapped_ids = [group_key_map.get(gid, gid) for gid in raw_group_ids]
        if raw_group_ids:
            ref.groups.clear()
        if mapped_ids:
            for gid in mapped_ids:
                g = db.query(Group).filter(Group.user_id == user.id, Group.group_key == gid).first()
                if g:
                    ref.groups.append(g)
        if not ref.groups:
            ref.groups.append(ungrouped)

    db.commit()

    # --- Import notes ---
    _import_notes(user.id, data.get("notes", {}), ref_key_map, db)

    # --- Import tags ---
    tag_map = _import_tags(user.id, data.get("note_tags", []), db)

    # --- Import standalone notes ---
    _import_standalone_notes(user.id, data.get("standalone_notes", []), tag_map, db)

    # --- Restore images from ZIP (sanitized against path traversal) ---
    os.makedirs(IMAGES_DIR, exist_ok=True)
    for name in zf.namelist():
        if name.startswith("images/") and len(name) > 7:
            img_name = os.path.basename(name[7:])
            if not img_name:
                continue
            img_path = os.path.join(IMAGES_DIR, img_name)
            if not os.path.exists(img_path):
                with open(img_path, "wb") as f:
                    f.write(zf.read(name))

    # --- Import daily plans ---
    _import_daily_plans(user.id, data.get("daily_plans", []), db)

    zf.close()
    return {"success": True}
